#!/usr/bin/env python
"""
Streamlit Web UI for the Document Reconstruction Pipeline.

Run with:
    streamlit run src/app.py

Features:
- Upload image files (PNG, JPG, TIFF, BMP)
- Real-time processing with progress indicator
- Preview of extracted text, equations, and tables
- Download results in multiple formats
"""

import sys
import os
from pathlib import Path

# Suppress verbose warnings from dependencies
os.environ["TF_CPP_MIN_LOG_LEVEL"] = "3"
os.environ["PYTHONWARNINGS"] = "ignore"
os.environ["NO_ALBUMENTATIONS_UPDATE"] = "1"
os.environ["DISABLE_MODEL_SOURCE_CHECK"] = "True"  # Suppress Pix2Text model check messages

import warnings
warnings.filterwarnings("ignore", category=UserWarning)
warnings.filterwarnings("ignore", category=FutureWarning)

# Add src directory to path for imports when running as script
_src_dir = Path(__file__).parent
if str(_src_dir) not in sys.path:
    sys.path.insert(0, str(_src_dir))

import streamlit as st
import tempfile
import json
import time
from typing import Optional
import io
import numpy as np

import logging
logging.getLogger('ppocr').setLevel(logging.ERROR)
logging.getLogger('pix2tex').setLevel(logging.ERROR)


class NumpyJSONEncoder(json.JSONEncoder):
    """JSON encoder that handles numpy types."""
    def default(self, obj):
        if isinstance(obj, np.ndarray):
            return obj.tolist()
        if isinstance(obj, np.integer):
            return int(obj)
        if isinstance(obj, np.floating):
            return float(obj)
        if isinstance(obj, np.bool_):
            return bool(obj)
        return super().default(obj)


def convert_numpy_types(obj):
    """Recursively convert numpy types to Python native types."""
    if isinstance(obj, dict):
        return {k: convert_numpy_types(v) for k, v in obj.items()}
    elif isinstance(obj, list):
        return [convert_numpy_types(item) for item in obj]
    elif isinstance(obj, np.ndarray):
        return obj.tolist()
    elif isinstance(obj, np.integer):
        return int(obj)
    elif isinstance(obj, np.floating):
        return float(obj)
    elif isinstance(obj, np.bool_):
        return bool(obj)
    return obj


# Page config must be first Streamlit command
st.set_page_config(
    page_title="Document Reconstruction",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded"
)


def load_css():
    """Load custom CSS styles."""
    st.markdown("""
    <style>
    .main-header {
        font-size: 2.5rem;
        font-weight: bold;
        color: #1E88E5;
        text-align: center;
        margin-bottom: 1rem;
    }
    .sub-header {
        font-size: 1.2rem;
        color: #888;
        text-align: center;
        margin-bottom: 2rem;
    }
    .block-preview {
        background-color: rgba(30, 136, 229, 0.1);
        border-radius: 5px;
        padding: 1rem;
        margin: 0.5rem 0;
        border-left: 4px solid #1E88E5;
        color: inherit;
    }
    .block-preview p {
        color: inherit;
        margin: 0.5rem 0;
    }
    .block-preview small {
        color: #888;
    }
    .equation-block {
        border-left-color: #9C27B0;
        background-color: rgba(156, 39, 176, 0.1);
    }
    .table-block {
        border-left-color: #4CAF50;
        background-color: rgba(76, 175, 80, 0.1);
    }
    </style>
    """, unsafe_allow_html=True)


def init_session_state():
    """Initialize session state variables."""
    if "document" not in st.session_state:
        st.session_state.document = None
    if "processing" not in st.session_state:
        st.session_state.processing = False
    if "error" not in st.session_state:
        st.session_state.error = None


@st.cache_data(ttl=300)  # Cache for 5 minutes
def check_engine_availability():
    """Check which OCR engines are available."""
    engines = {}
    
    # Check Tesseract
    try:
        import pytesseract
        pytesseract.get_tesseract_version()
        engines["tesseract"] = {"available": True, "error": None}
    except Exception as e:
        engines["tesseract"] = {"available": False, "error": str(e)[:50]}
    
    # Check EasyOCR
    try:
        import easyocr
        engines["easyocr"] = {"available": True, "error": None}
    except ImportError:
        engines["easyocr"] = {"available": False, "error": "pip install easyocr"}
    
    # Check PaddleOCR
    try:
        from paddleocr import PaddleOCR
        engines["paddleocr"] = {"available": True, "error": None}
    except ImportError:
        engines["paddleocr"] = {"available": False, "error": "pip install paddleocr"}
    except Exception as e:
        engines["paddleocr"] = {"available": False, "error": str(e)[:50]}
    
    # Check pix2tex (Math OCR)
    try:
        from pix2tex.cli import LatexOCR
        try:
            _ = LatexOCR()
            engines["pix2tex"] = {"available": True, "error": None}
        except (ValueError, RuntimeError) as e:
            if "std_range" in str(e) or "albumentations" in str(e).lower():
                engines["pix2tex"] = {"available": False, "error": "pip install 'albumentations<1.4.0'"}
            else:
                engines["pix2tex"] = {"available": False, "error": str(e)[:50]}
    except ImportError:
        engines["pix2tex"] = {"available": False, "error": "pip install pix2tex"}
    except Exception as e:
        engines["pix2tex"] = {"available": False, "error": str(e)[:50]}
    
    # Check Pix2Text (layout detection)
    try:
        from pix2text import Pix2Text
        engines["pix2text"] = {"available": True, "error": None}
    except ImportError:
        engines["pix2text"] = {"available": False, "error": "pip install pix2text"}
    except Exception as e:
        engines["pix2text"] = {"available": False, "error": str(e)[:50]}
    
    # Check LayoutParser
    try:
        import layoutparser
        engines["layoutparser"] = {"available": True, "error": None}
    except ImportError:
        engines["layoutparser"] = {"available": False, "error": "Using classical CV fallback"}
    
    return engines


def render_engine_status_row(name: str, available: bool, error: str = None):
    """Render a single engine status row with colored indicator."""
    col1, col2 = st.columns([4, 1])
    with col1:
        st.markdown(f"**{name}**")
        if not available and error:
            st.caption(f"{error}")
    with col2:
        if available:
            st.markdown("<div style='text-align:right'><span style='color:#22c55e;font-size:20px'>●</span></div>", unsafe_allow_html=True)
        else:
            st.markdown("<div style='text-align:right'><span style='color:#ef4444;font-size:20px'>●</span></div>", unsafe_allow_html=True)


def render_sidebar():
    """Render sidebar with settings."""
    st.sidebar.header("⚙️ Settings")
    
    # Check engine availability
    engines = check_engine_availability()
    
    # Processing options
    st.sidebar.subheader("Processing")
    
    use_gpu = st.sidebar.checkbox(
        "Use GPU",
        value=False,
        help="Enable GPU acceleration if available"
    )
    
    debug_mode = st.sidebar.checkbox(
        "Debug Mode",
        value=False,
        help="Save debug images with bounding boxes"
    )
    
    # DPI setting (hidden - used internally if PDF support added later)
    dpi = 400
    
    # Layout detection options
    st.sidebar.subheader("Layout Detection")
    
    # Display names mapping
    layout_options = {
        "Default": "default",
        "Classical CV": "classical", 
        "LayoutParser": "layoutparser",
        "PaddleOCR": "paddleocr",
        "Pix2Text": "pix2text"
    }
    layout_display = st.sidebar.selectbox(
        "Method",
        list(layout_options.keys()),
        index=0,
        help="Method for detecting document structure (Default = tries best available)"
    )
    layout_method = layout_options[layout_display]
    
    # Only show warning if selected method is unavailable
    if layout_method == "pix2text" and not engines.get("pix2text", {}).get("available"):
        st.sidebar.warning("Pix2Text unavailable, will fall back to Classical CV")
    elif layout_method == "layoutparser" and not engines.get("layoutparser", {}).get("available"):
        st.sidebar.warning("LayoutParser unavailable, will fall back to Classical CV")
    elif layout_method == "paddleocr" and not engines.get("paddleocr", {}).get("available"):
        st.sidebar.warning("PaddleOCR unavailable, will fall back to Classical CV")
    
    # OCR options
    st.sidebar.subheader("OCR Engine")
    
    ocr_options = {
        "Tesseract OCR": "tesseract",
        "EasyOCR": "easyocr",
        "PaddleOCR": "paddleocr"
    }
    ocr_display = st.sidebar.selectbox(
        "Text OCR",
        list(ocr_options.keys()),
        index=0,
        help="Primary OCR engine for text extraction"
    )
    ocr_engine = ocr_options[ocr_display]
    
    # Only show warning if selected engine is unavailable
    if ocr_engine in engines and not engines[ocr_engine]["available"]:
        st.sidebar.warning(f"{ocr_display} unavailable, will fall back to Tesseract OCR")
    
    math_options = {
        "pix2tex (LaTeX OCR)": "pix2tex",
        "Simple (Text-based)": "simple"
    }
    math_display = st.sidebar.selectbox(
        "Math OCR",
        list(math_options.keys()),
        index=0,
        help="Engine for equation recognition"
    )
    math_engine = math_options[math_display]
    
    # Only show warning if selected engine is unavailable
    if math_engine in engines and not engines[math_engine]["available"]:
        st.sidebar.warning("pix2tex unavailable, will fall back to Simple recognition")
    
    confidence_threshold = st.sidebar.slider(
        "Confidence Threshold",
        min_value=0.0,
        max_value=1.0,
        value=0.65,
        step=0.05,
        help="Minimum confidence for OCR results"
    )
    
    # Engine Status Overview
    with st.sidebar.expander("Engine Status", expanded=False):
        st.markdown("**Layout Detection**")
        st.divider()
        render_engine_status_row("Pix2Text", engines.get("pix2text", {}).get("available", False), engines.get("pix2text", {}).get("error"))
        st.divider()
        render_engine_status_row("LayoutParser", engines.get("layoutparser", {}).get("available", False), engines.get("layoutparser", {}).get("error"))
        st.divider()
        
        st.markdown("**Text OCR**")
        st.divider()
        for eng in ["Tesseract", "EasyOCR", "PaddleOCR"]:
            eng_key = eng.lower()
            render_engine_status_row(eng, engines.get(eng_key, {}).get("available", False), engines.get(eng_key, {}).get("error"))
            st.divider()
        
        st.markdown("**Math OCR**")
        st.divider()
        render_engine_status_row("pix2tex", engines.get("pix2tex", {}).get("available", False), engines.get("pix2tex", {}).get("error"))
    
    # Export options
    st.sidebar.subheader("Export")
    
    export_formats = st.sidebar.multiselect(
        "Formats",
        ["JSON", "Markdown", "DOCX", "LaTeX"],
        default=["JSON", "Markdown"],
        help="Output formats to generate"
    )
    
    return {
        "use_gpu": use_gpu,
        "debug_mode": debug_mode,
        "dpi": dpi,
        "layout_method": layout_method,
        "ocr_engine": ocr_engine,
        "math_engine": math_engine,
        "confidence_threshold": confidence_threshold,
        "export_formats": [f.lower() for f in export_formats]
    }


def process_document(uploaded_file, settings) -> Optional[dict]:
    """Process the uploaded document."""
    try:
        # Import here to avoid slow initial load
        from utils.io import load_pdf, load_image, detect_input_type
        from utils.assembler import DocumentAssembler
        from utils.export import DocumentExporter
        import numpy as np
        import cv2
        
        # Create temp directory for processing
        with tempfile.TemporaryDirectory(prefix="doc_recon_") as temp_dir:
            temp_dir = Path(temp_dir)
            
            # Save uploaded file
            input_path = temp_dir / uploaded_file.name
            with open(input_path, "wb") as f:
                f.write(uploaded_file.getbuffer())
            
            # Detect type and load images
            input_type = detect_input_type(input_path)
            
            progress_bar = st.progress(0, text="Loading document...")
            
            if input_type == "pdf":
                images = load_pdf(input_path, dpi=settings["dpi"])
                pdf_path = str(input_path)
            elif input_type == "image":
                images = [load_image(input_path)]
                pdf_path = None
            else:
                raise ValueError(f"Unsupported file type: {uploaded_file.name}")
            
            progress_bar.progress(20, text=f"Loaded {len(images)} page(s)")
            
            # Create assembler
            output_dir = temp_dir / "output"
            output_dir.mkdir()
            
            # Check if selected engine is available, otherwise fall back
            actual_ocr_engine = settings["ocr_engine"]
            actual_math_engine = settings["math_engine"]
            
            # Check OCR engine availability
            if actual_ocr_engine == "easyocr":
                try:
                    import easyocr
                except ImportError:
                    st.warning(f"⚠️ {actual_ocr_engine} not available, falling back to tesseract")
                    actual_ocr_engine = "tesseract"
            elif actual_ocr_engine == "paddleocr":
                try:
                    from paddleocr import PaddleOCR
                except ImportError:
                    st.warning(f"⚠️ {actual_ocr_engine} not available, falling back to tesseract")
                    actual_ocr_engine = "tesseract"
            
            # Check math engine availability
            if actual_math_engine == "pix2tex":
                try:
                    from pix2tex.cli import LatexOCR
                except ImportError:
                    st.warning(f"⚠️ {actual_math_engine} not available, falling back to simple")
                    actual_math_engine = "simple"
            
            layout_method = settings.get("layout_method", "default")
            st.info(f"🔧 Using: **{layout_method}** layout, **{actual_ocr_engine}** OCR, **{actual_math_engine}** math")
            
            assembler = DocumentAssembler(
                use_gpu=settings["use_gpu"],
                debug_mode=settings["debug_mode"],
                output_dir=output_dir,
                ocr_engine=actual_ocr_engine,
                math_engine=actual_math_engine,
                layout_method=layout_method
            )
            
            # Process document with progress updates
            total_pages = len(images)
            pages = []
            preprocessed_images = []
            
            for i, image in enumerate(images):
                progress_bar.progress(
                    20 + int(60 * (i / total_pages)),
                    text=f"Processing page {i+1}/{total_pages}..."
                )
                
                # Store preprocessed image for preview
                from utils.images import preprocess_image
                preprocess_result = preprocess_image(
                    image,
                    deskew_enabled=True,
                    denoise_enabled=True,
                    enhance_contrast_enabled=True
                )
                preprocessed_images.append({
                    "image": preprocess_result.image,
                    "original": image,
                    "deskew_angle": preprocess_result.deskew_angle,
                    "transformations": preprocess_result.transformations
                })
                
                page = assembler.process_page(
                    image,
                    page_number=i+1,
                    pdf_path=pdf_path
                )
                pages.append(page)
            
            # Store preprocessed images in session state
            st.session_state.preprocessed_images = preprocessed_images
            
            progress_bar.progress(80, text="Assembling document...")
            
            # Assemble document
            from utils.assembler import Document, DocumentMetrics
            import uuid
            from datetime import datetime
            
            document = Document(
                task_id=str(uuid.uuid4()),
                source_file=uploaded_file.name,
                pages=pages
            )
            
            # Calculate metrics
            document.metrics = assembler._calculate_metrics(document, 0)
            document.markdown = assembler._generate_markdown(document)
            document.docx_manifest = assembler._generate_docx_manifest(document)
            
            # Extract metadata
            if pages:
                assembler._extract_document_metadata(document, pages[0])
            
            progress_bar.progress(100, text="Complete!")
            time.sleep(0.5)
            progress_bar.empty()
            
            return document.to_dict()
            
    except Exception as e:
        st.error(f"Processing error: {str(e)}")
        import traceback
        if settings.get("debug_mode"):
            st.code(traceback.format_exc())
        return None


def check_engine_fallbacks(document: dict) -> list:
    """Check if any blocks used fallback engines and return warnings."""
    fallbacks = []
    
    for page in document.get("pages", []):
        for block in page.get("blocks", []):
            ocr_meta = block.get("metadata", {}).get("ocr", {})
            if ocr_meta.get("used_secondary"):
                fallback_info = {
                    "block_type": block.get("type", "unknown"),
                    "primary_engine": ocr_meta.get("primary_engine", "unknown"),
                    "secondary_engine": ocr_meta.get("secondary_engine", "unknown"),
                    "primary_confidence": ocr_meta.get("primary_confidence", 0),
                    "final_confidence": block.get("confidence", 0),
                    "threshold": ocr_meta.get("confidence_threshold", 0.65),
                    "reason": ocr_meta.get("reason", "")
                }
                fallbacks.append(fallback_info)
    
    return fallbacks


def render_engine_fallback_notices(document: dict):
    """Display notices about engine fallbacks if any occurred."""
    fallbacks = check_engine_fallbacks(document)
    
    if fallbacks:
        with st.expander(f"⚠️ Engine Fallback Notices ({len(fallbacks)} blocks)", expanded=False):
            st.warning(
                f"**{len(fallbacks)} block(s)** used a fallback OCR engine because the primary engine "
                f"did not meet the confidence threshold."
            )
            
            for i, fb in enumerate(fallbacks, 1):
                st.markdown(
                    f"**Block {i}** ({fb['block_type']}): "
                    f"`{fb['primary_engine']}` → `{fb['secondary_engine']}` | "
                    f"Primary confidence: {fb['primary_confidence']:.0%} (threshold: {fb['threshold']:.0%}) | "
                    f"Final confidence: {fb['final_confidence']:.0%}"
                )


def render_metrics(metrics: dict):
    """Render document metrics."""
    cols = st.columns(5)
    
    with cols[0]:
        conf = metrics.get("global_confidence", 0)
        st.metric("Confidence", f"{conf:.0%}")
    
    with cols[1]:
        cov = metrics.get("coverage_pct", 0)
        st.metric("Coverage", f"{cov:.0f}%")
    
    with cols[2]:
        text_blocks = metrics.get("text_blocks", {})
        st.metric("Text Blocks", text_blocks.get("total", 0))
    
    with cols[3]:
        equations = metrics.get("equations", {})
        st.metric("Equations", equations.get("total", 0))
    
    with cols[4]:
        tables = metrics.get("tables", {})
        st.metric("Tables", tables.get("total", 0))


def render_page_content(page: dict):
    """Render content of a single page."""
    blocks = sorted(page.get("blocks", []), key=lambda b: b.get("reading_order", 0))
    
    if not blocks:
        st.info("No blocks detected on this page.")
        return
    
    for block in blocks:
        block_type = block.get("type", "unknown")
        confidence = block.get("confidence", 0)
        
        if block_type in ["equation_block", "equation_inline"]:
            latex = block.get("latex", "")
            if latex:
                with st.container():
                    st.caption(f"🔢 **Equation** (confidence: {confidence:.0%})")
                    try:
                        st.latex(latex)
                    except:
                        st.code(latex, language="latex")
        
        elif block_type == "table":
            table = block.get("table", {})
            if table:
                with st.container():
                    st.caption(f"📋 **Table** (confidence: {confidence:.0%})")
                    rows = table.get("struct", {}).get("rows", [])
                    if rows and len(rows) > 1:
                        import pandas as pd
                        try:
                            df = pd.DataFrame(rows[1:], columns=rows[0])
                            st.dataframe(df, use_container_width=True)
                        except:
                            st.text(table.get("markdown", "Table could not be displayed"))
                    elif table.get("markdown"):
                        st.markdown(table.get("markdown"))
        
        elif block.get("text"):
            text = block.get("text", "")
            
            if block_type == "title":
                st.markdown(f"# {text}")
            elif block_type == "heading":
                st.markdown(f"## {text}")
            else:
                with st.expander(f"📝 {block_type.title()} (confidence: {confidence:.0%})", expanded=True):
                    st.write(text)
        
        elif block.get("figure_path"):
            st.caption(f"🖼️ **Figure** (confidence: {confidence:.0%})")
            st.image(block.get("figure_path"), use_container_width=True)


def render_bounding_box_preview(document: dict, image_bytes: bytes):
    """Render the original image with detected bounding boxes overlaid."""
    import cv2
    import numpy as np
    from PIL import Image
    import io
    
    # Color map for different block types
    BLOCK_COLORS = {
        "title": (255, 0, 0),       # Blue
        "heading": (255, 128, 0),   # Orange-blue
        "paragraph": (0, 255, 0),   # Green
        "text": (0, 255, 0),        # Green
        "equation_block": (0, 0, 255),   # Red
        "equation_inline": (128, 0, 255), # Purple
        "table": (255, 255, 0),     # Cyan
        "figure": (255, 0, 255),    # Magenta
        "caption": (128, 128, 0),   # Teal
        "list": (0, 128, 255),      # Orange
        "header": (128, 128, 128),  # Gray
        "footer": (128, 128, 128),  # Gray
        "unknown": (200, 200, 200), # Light gray
    }
    
    pages = document.get("pages", [])
    if not pages:
        st.info("No pages to display")
        return
    
    # Load the original image from bytes
    try:
        pil_image = Image.open(io.BytesIO(image_bytes))
        image = cv2.cvtColor(np.array(pil_image), cv2.COLOR_RGB2BGR)
    except Exception as e:
        st.error(f"Could not load image for preview: {e}")
        return
    
    # Page selector for multi-page (though images are single page)
    page_idx = 0
    if len(pages) > 1:
        page_idx = st.selectbox(
            "Select page for preview",
            range(len(pages)),
            format_func=lambda x: f"Page {x + 1}",
            key="bbox_page_select"
        ) or 0
    
    page = pages[page_idx]
    blocks = page.get("blocks", [])
    
    if not blocks:
        st.info("No blocks detected on this page")
        st.image(pil_image, caption="Original Image", use_container_width=True)
        return
    
    # Filter options
    st.markdown("**Filter block types:**")
    col1, col2, col3, col4 = st.columns(4)
    
    all_types = list(set(b.get("type", "unknown") for b in blocks))
    selected_types = []
    
    with col1:
        if st.checkbox("Text/Paragraphs", value=True, key="show_text"):
            selected_types.extend(["paragraph", "text", "title", "heading"])
    with col2:
        if st.checkbox("Equations", value=True, key="show_eq"):
            selected_types.extend(["equation_block", "equation_inline"])
    with col3:
        if st.checkbox("Tables", value=True, key="show_table"):
            selected_types.append("table")
    with col4:
        if st.checkbox("Other", value=True, key="show_other"):
            selected_types.extend(["figure", "caption", "list", "header", "footer", "unknown"])
    
    # Draw bounding boxes
    preview_image = image.copy()
    
    for block in blocks:
        block_type = block.get("type", "unknown")
        if block_type not in selected_types:
            continue
            
        bbox = block.get("bbox", {})
        if not bbox:
            continue
        
        x1 = int(bbox.get("x1", 0))
        y1 = int(bbox.get("y1", 0))
        x2 = int(bbox.get("x2", 0))
        y2 = int(bbox.get("y2", 0))
        
        if x2 <= x1 or y2 <= y1:
            continue
        
        color = BLOCK_COLORS.get(block_type, (200, 200, 200))
        confidence = block.get("confidence", 0)
        
        # Draw rectangle
        cv2.rectangle(preview_image, (x1, y1), (x2, y2), color, 2)
        
        # Draw label background
        label = f"{block_type} ({confidence:.0%})"
        (label_w, label_h), baseline = cv2.getTextSize(label, cv2.FONT_HERSHEY_SIMPLEX, 0.5, 1)
        cv2.rectangle(preview_image, (x1, y1 - label_h - 8), (x1 + label_w + 4, y1), color, -1)
        
        # Draw label text
        cv2.putText(preview_image, label, (x1 + 2, y1 - 4), 
                    cv2.FONT_HERSHEY_SIMPLEX, 0.5, (255, 255, 255), 1)
    
    # Convert back to RGB for display
    preview_rgb = cv2.cvtColor(preview_image, cv2.COLOR_BGR2RGB)
    
    # Display
    st.image(preview_rgb, caption="Detected Bounding Boxes", use_container_width=True)
    
    # Legend
    st.markdown("**Legend:**")
    legend_cols = st.columns(4)
    type_list = list(BLOCK_COLORS.keys())
    for i, block_type in enumerate(type_list[:8]):  # Show first 8
        color = BLOCK_COLORS[block_type]
        # Convert BGR to hex
        hex_color = "#{:02x}{:02x}{:02x}".format(color[2], color[1], color[0])
        with legend_cols[i % 4]:
            st.markdown(f"<span style='color:{hex_color}'>■</span> {block_type}", unsafe_allow_html=True)
    
    # Block count summary
    block_counts = {}
    for b in blocks:
        t = b.get("type", "unknown")
        block_counts[t] = block_counts.get(t, 0) + 1
    
    st.markdown("**Detected blocks:**")
    st.write(" | ".join([f"{t}: {c}" for t, c in sorted(block_counts.items())]))


def render_preprocessed_preview():
    """Render side-by-side comparison of original vs preprocessed image."""
    import cv2
    import numpy as np
    
    preprocessed_data = st.session_state.get("preprocessed_images", [])
    
    if not preprocessed_data:
        st.info("No preprocessed images available. Process a document first.")
        return
    
    # Page selector
    page_idx = 0
    if len(preprocessed_data) > 1:
        page_idx = st.selectbox(
            "Select page",
            range(len(preprocessed_data)),
            format_func=lambda x: f"Page {x + 1}",
            key="preprocess_page_select"
        ) or 0
    
    data = preprocessed_data[page_idx]
    original = data["original"]
    preprocessed = data["image"]
    deskew_angle = data.get("deskew_angle", 0)
    transformations = data.get("transformations", [])
    
    # Display info
    st.markdown("**Preprocessing Applied:**")
    if transformations:
        st.write(" → ".join(transformations))
    else:
        st.write("No transformations applied (clean image detected)")
    
    if abs(deskew_angle) > 0.1:
        st.write(f"Deskew angle: **{deskew_angle:.2f}°**")
    
    # Convert images to RGB for display
    if len(original.shape) == 3:
        original_rgb = cv2.cvtColor(original, cv2.COLOR_BGR2RGB)
    else:
        original_rgb = cv2.cvtColor(original, cv2.COLOR_GRAY2RGB)
    
    if len(preprocessed.shape) == 3:
        preprocessed_rgb = cv2.cvtColor(preprocessed, cv2.COLOR_BGR2RGB)
    else:
        preprocessed_rgb = cv2.cvtColor(preprocessed, cv2.COLOR_GRAY2RGB)
    
    # Display side by side
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("**Original Image**")
        st.image(original_rgb, use_container_width=True)
        st.caption(f"Size: {original.shape[1]}×{original.shape[0]}")
    
    with col2:
        st.markdown("**Preprocessed Image**")
        st.image(preprocessed_rgb, use_container_width=True)
        st.caption(f"Size: {preprocessed.shape[1]}×{preprocessed.shape[0]}")
    
    # Image statistics
    with st.expander("Image Statistics"):
        col1, col2 = st.columns(2)
        
        orig_gray = cv2.cvtColor(original, cv2.COLOR_BGR2GRAY) if len(original.shape) == 3 else original
        prep_gray = cv2.cvtColor(preprocessed, cv2.COLOR_BGR2GRAY) if len(preprocessed.shape) == 3 else preprocessed
        
        with col1:
            st.markdown("**Original:**")
            st.write(f"- Mean intensity: {np.mean(orig_gray):.1f}")
            st.write(f"- Std deviation: {np.std(orig_gray):.1f}")
            laplacian_var = cv2.Laplacian(orig_gray, cv2.CV_64F).var()
            st.write(f"- Sharpness (Laplacian): {laplacian_var:.1f}")
        
        with col2:
            st.markdown("**Preprocessed:**")
            st.write(f"- Mean intensity: {np.mean(prep_gray):.1f}")
            st.write(f"- Std deviation: {np.std(prep_gray):.1f}")
            laplacian_var = cv2.Laplacian(prep_gray, cv2.CV_64F).var()
            st.write(f"- Sharpness (Laplacian): {laplacian_var:.1f}")


def generate_docx_bytes(document: dict) -> bytes:
    """Generate a DOCX file from document and return as bytes."""
    try:
        from docx import Document as DocxDocument
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = DocxDocument()
        
        # Title
        title = document.get("metadata", {}).get("title")
        if title and title != "UNREADABLE":
            heading = doc.add_heading(title, 0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Authors
        authors = document.get("metadata", {}).get("authors")
        if authors and authors != "UNREADABLE":
            p = doc.add_paragraph(authors)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Abstract
        abstract = document.get("metadata", {}).get("abstract")
        if abstract:
            doc.add_heading("Abstract", level=1)
            doc.add_paragraph(abstract)
        
        # Process pages
        for page in document.get("pages", []):
            blocks = sorted(page.get("blocks", []), key=lambda b: b.get("reading_order", 0))
            
            for block in blocks:
                block_type = block.get("type", "")
                
                if block_type == "title":
                    continue  # Already added
                elif block_type == "heading":
                    if block.get("text"):
                        doc.add_heading(block["text"], level=2)
                elif block.get("text"):
                    doc.add_paragraph(block["text"])
                elif block.get("latex"):
                    # Add equation as formatted text
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run(f"[Equation: {block['latex']}]")
                    run.italic = True
                elif block.get("table"):
                    # Add table
                    table_data = block["table"]
                    rows = table_data.get("struct", {}).get("rows", [])
                    if rows and len(rows) > 0:
                        num_rows = len(rows)
                        num_cols = len(rows[0]) if rows else 0
                        if num_cols > 0:
                            table = doc.add_table(rows=num_rows, cols=num_cols)
                            table.style = 'Table Grid'
                            for i, row_data in enumerate(rows):
                                row = table.rows[i]
                                for j, cell_text in enumerate(row_data):
                                    if j < len(row.cells):
                                        row.cells[j].text = str(cell_text)
        
        # Save to bytes
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        return file_stream.getvalue()
        
    except ImportError:
        return None
    except Exception as e:
        st.error(f"Error generating DOCX: {e}")
        return None


def generate_pdf_bytes(document: dict) -> bytes:
    """Generate a PDF from document markdown using fpdf2."""
    markdown = document.get("markdown", "")
    if not markdown:
        return None
    
    try:
        from fpdf import FPDF
        from fpdf.enums import XPos, YPos
        
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        
        # Parse markdown and render to PDF
        lines = markdown.split('\n')
        in_table = False
        table_rows = []
        in_code_block = False
        
        for line in lines:
            line = line.rstrip()
            
            # Code blocks
            if line.startswith('```'):
                in_code_block = not in_code_block
                continue
            
            if in_code_block:
                pdf.set_font('Courier', '', 10)
                pdf.set_fill_color(240, 240, 240)
                pdf.multi_cell(0, 5, line, fill=True, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
                continue
            
            # Table handling
            if '|' in line and not line.startswith('|--'):
                cells = [c.strip() for c in line.split('|')[1:-1]]
                if cells:
                    if not in_table:
                        in_table = True
                        table_rows = []
                    table_rows.append(cells)
                continue
            elif line.startswith('|--') or (line.startswith('|') and '-' in line):
                continue  # Skip separator rows
            elif in_table and table_rows:
                # Render table
                _render_pdf_table(pdf, table_rows)
                table_rows = []
                in_table = False
            
            # Headers
            if line.startswith('# '):
                pdf.set_font('Helvetica', 'B', 18)
                pdf.multi_cell(0, 10, line[2:], new_x=XPos.LMARGIN, new_y=YPos.NEXT)
                pdf.ln(3)
            elif line.startswith('## '):
                pdf.set_font('Helvetica', 'B', 14)
                pdf.multi_cell(0, 8, line[3:], new_x=XPos.LMARGIN, new_y=YPos.NEXT)
                pdf.ln(2)
            elif line.startswith('### '):
                pdf.set_font('Helvetica', 'B', 12)
                pdf.multi_cell(0, 7, line[4:], new_x=XPos.LMARGIN, new_y=YPos.NEXT)
                pdf.ln(2)
            elif line.strip() == '':
                pdf.ln(3)
            elif line.startswith('- ') or line.startswith('* '):
                pdf.set_font('Helvetica', '', 11)
                pdf.cell(5, 6, chr(8226))  # Bullet
                pdf.multi_cell(0, 6, line[2:], new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            else:
                pdf.set_font('Helvetica', '', 11)
                # Handle bold/italic
                clean_line = line.replace('**', '').replace('*', '').replace('`', '')
                pdf.multi_cell(0, 6, clean_line, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        
        # Render any remaining table
        if in_table and table_rows:
            _render_pdf_table(pdf, table_rows)
        
        return bytes(pdf.output())
        
    except ImportError:
        pass
    except Exception as e:
        import logging
        logging.getLogger(__name__).error(f"PDF generation error: {e}")
    
    return None


def _render_pdf_table(pdf, rows):
    """Render a table to PDF."""
    from fpdf.enums import XPos, YPos
    
    if not rows:
        return
    
    num_cols = max(len(row) for row in rows)
    page_width = pdf.w - pdf.l_margin - pdf.r_margin
    col_width = page_width / num_cols
    
    for i, row in enumerate(rows):
        # Header row gets bold
        if i == 0:
            pdf.set_font('Helvetica', 'B', 10)
            pdf.set_fill_color(230, 230, 230)
        else:
            pdf.set_font('Helvetica', '', 10)
            pdf.set_fill_color(255, 255, 255)
        
        for j, cell in enumerate(row):
            fill = (i == 0)
            pdf.cell(col_width, 7, str(cell)[:30], border=1, fill=fill)
        
        pdf.ln()
    
    pdf.ln(3)


def render_downloads(document: dict, settings: dict):
    """Render download buttons."""
    st.subheader("📥 Downloads")
    
    cols = st.columns(5)
    
    # JSON
    with cols[0]:
        json_str = json.dumps(document, indent=2, cls=NumpyJSONEncoder)
        st.download_button(
            "📄 JSON",
            json_str,
            file_name="document.json",
            mime="application/json",
            use_container_width=True
        )
    
    # Markdown
    with cols[1]:
        markdown = document.get("markdown", "")
        st.download_button(
            "📝 Markdown",
            markdown,
            file_name="document.md",
            mime="text/markdown",
            use_container_width=True
        )
    
    # DOCX - actual DOCX generation
    with cols[2]:
        docx_bytes = generate_docx_bytes(document)
        if docx_bytes:
            st.download_button(
                "📋 DOCX",
                docx_bytes,
                file_name="document.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
        else:
            st.button(
                "📋 DOCX ❌",
                use_container_width=True,
                help="Install python-docx: pip install python-docx",
                disabled=True
            )
    
    # LaTeX
    with cols[3]:
        latex_content = generate_simple_latex(document)
        st.download_button(
            "📐 LaTeX",
            latex_content,
            file_name="document.tex",
            mime="text/x-tex",
            use_container_width=True
        )
    
    # PDF
    with cols[4]:
        pdf_bytes = generate_pdf_bytes(document)
        if pdf_bytes:
            st.download_button(
                "📕 PDF",
                pdf_bytes,
                file_name="document.pdf",
                mime="application/pdf",
                use_container_width=True
            )
        else:
            st.button(
                "📕 PDF ❌",
                use_container_width=True,
                help="Install: pip install fpdf2",
                disabled=True
            )


def generate_simple_latex(document: dict) -> str:
    """Generate simple LaTeX from document."""
    lines = [
        r"\documentclass[11pt]{article}",
        r"\usepackage{amsmath,amssymb,graphicx}",
        r"\begin{document}",
        ""
    ]
    
    # Title
    title = document.get("metadata", {}).get("title")
    if title and title != "UNREADABLE":
        lines.append(r"\title{" + title.replace("_", r"\_") + "}")
        lines.append(r"\maketitle")
        lines.append("")
    
    # Content from markdown (simplified)
    markdown = document.get("markdown", "")
    for line in markdown.split("\n"):
        if line.startswith("# "):
            lines.append(r"\section*{" + line[2:] + "}")
        elif line.startswith("## "):
            lines.append(r"\subsection*{" + line[3:] + "}")
        elif line.startswith("$$"):
            lines.append(line)
        elif line.strip():
            lines.append(line.replace("_", r"\_").replace("&", r"\&"))
        lines.append("")
    
    lines.append(r"\end{document}")
    
    return "\n".join(lines)


def main():
    """Main application."""
    load_css()
    init_session_state()
    
    # Header
    st.markdown('<h1 class="main-header">📄 Document Reconstruction</h1>', unsafe_allow_html=True)
    st.markdown(
        '<p class="sub-header">Convert scanned documents to structured formats with AI-powered OCR</p>',
        unsafe_allow_html=True
    )
    
    # Sidebar settings
    settings = render_sidebar()
    
    # File upload
    st.markdown("---")
    
    uploaded_file = st.file_uploader(
        "Upload a document image",
        type=["png", "jpg", "jpeg", "tiff", "bmp"],
        help="Upload an image file to process (PNG, JPG, TIFF, BMP)"
    )
    
    if uploaded_file:
        col1, col2 = st.columns([2, 1])
        
        with col1:
            st.info(f"📁 **{uploaded_file.name}** ({uploaded_file.size / 1024:.1f} KB)")
        
        with col2:
            process_btn = st.button(
                "🚀 Process Document",
                use_container_width=True,
                type="primary"
            )
        
        if process_btn:
            with st.spinner("Processing document..."):
                result = process_document(uploaded_file, settings)
                
                if result:
                    # Convert numpy types to native Python types for JSON compatibility
                    st.session_state.document = convert_numpy_types(result)
                    # Store image bytes for bounding box preview
                    uploaded_file.seek(0)
                    st.session_state.uploaded_image_bytes = uploaded_file.read()
                    st.success("✅ Document processed successfully!")
    
    # Display results
    if st.session_state.document:
        doc = st.session_state.document
        
        st.markdown("---")
        
        # Metrics
        st.subheader("📊 Processing Metrics")
        render_metrics(doc.get("metrics", {}))
        
        # Engine fallback notices
        render_engine_fallback_notices(doc)
        
        st.markdown("---")
        
        # Content tabs
        tabs = st.tabs(["Content", "Bounding Boxes", "Preprocessed", "Markdown", "Raw JSON"])
        
        with tabs[0]:
            pages = doc.get("pages", [])
            if len(pages) > 1:
                page_num = st.selectbox(
                    "Select page",
                    range(1, len(pages) + 1),
                    format_func=lambda x: f"Page {x}"
                )
                render_page_content(pages[page_num - 1])
            elif pages:
                render_page_content(pages[0])
            else:
                st.info("No content extracted")
        
        with tabs[1]:
            # Bounding box preview
            if st.session_state.get("uploaded_image_bytes"):
                render_bounding_box_preview(doc, st.session_state.uploaded_image_bytes)
            else:
                st.info("Upload and process an image to see bounding box preview")
        
        with tabs[2]:
            # Preprocessed image preview
            render_preprocessed_preview()
        
        with tabs[3]:
            markdown = doc.get("markdown", "")
            if markdown:
                st.code(markdown, language="markdown")
            else:
                st.info("No markdown content")
        
        with tabs[4]:
            st.json(convert_numpy_types(doc))
        
        st.markdown("---")
        
        # Downloads
        render_downloads(doc, settings)
    
    # Footer
    st.markdown("---")
    st.markdown(
        """
        <div style="text-align: center; color: #666; font-size: 0.8rem;">
            Document Reconstruction Pipeline v1.0 | 
            Built with Streamlit, OpenCV, and PyTorch
        </div>
        """,
        unsafe_allow_html=True
    )


if __name__ == "__main__":
    main()


