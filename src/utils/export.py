"""
Export module for document reconstruction.

Provides:
- Markdown export
- DOCX export (using python-docx)
- LaTeX export
- PDF export (via LaTeX)
"""

import logging
import subprocess
import tempfile
import shutil
from pathlib import Path
from typing import Dict, Any, Optional, List, Union
from dataclasses import dataclass

logger = logging.getLogger(__name__)


# ============================================================================
# Markdown Exporter
# ============================================================================

class MarkdownExporter:
    """Export document to Markdown format."""
    
    def __init__(
        self,
        math_delimiter_inline: str = "$",
        math_delimiter_block: str = "$$",
        include_page_breaks: bool = True
    ):
        self.math_delimiter_inline = math_delimiter_inline
        self.math_delimiter_block = math_delimiter_block
        self.include_page_breaks = include_page_breaks
    
    def export(
        self,
        document: Any,
        output_path: Union[str, Path]
    ) -> Path:
        """
        Export document to Markdown file.
        
        Args:
            document: Document object
            output_path: Output file path
            
        Returns:
            Path to the generated Markdown file
        """
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        # Use pre-generated markdown if available
        if hasattr(document, 'markdown') and document.markdown:
            markdown = document.markdown
        else:
            markdown = self._generate_markdown(document)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(markdown)
        
        logger.info(f"Exported Markdown to: {output_path}")
        return output_path
    
    def _generate_markdown(self, document: Any) -> str:
        """Generate Markdown from document structure."""
        lines = []
        
        # Document header
        if hasattr(document, 'title') and document.title:
            lines.append(f"# {document.title}")
            lines.append("")
        
        if hasattr(document, 'authors') and document.authors:
            lines.append(f"**Authors:** {document.authors}")
            lines.append("")
        
        if hasattr(document, 'abstract') and document.abstract:
            lines.append("## Abstract")
            lines.append("")
            lines.append(document.abstract)
            lines.append("")
        
        # Process pages
        for page in document.pages:
            if self.include_page_breaks and len(document.pages) > 1:
                lines.append("")
                lines.append("---")
                lines.append(f"*Page {page.page_number}*")
                lines.append("")
            
            # Process blocks in reading order
            sorted_blocks = sorted(page.blocks, key=lambda b: b.reading_order)
            
            for block in sorted_blocks:
                md = self._block_to_markdown(block)
                if md:
                    lines.append(md)
                    lines.append("")
        
        return "\n".join(lines)
    
    def _block_to_markdown(self, block: Any) -> str:
        """Convert a block to Markdown."""
        block_type = block.block_type if isinstance(block.block_type, str) else block.block_type.value
        
        if block_type == "title":
            return f"# {block.text}" if block.text else ""
        
        elif block_type == "heading":
            return f"## {block.text}" if block.text else ""
        
        elif block_type in ["paragraph", "abstract"]:
            return block.text if block.text else ""
        
        elif block_type == "equation_block" and block.latex:
            return f"{self.math_delimiter_block}\n{block.latex}\n{self.math_delimiter_block}"
        
        elif block_type == "equation_inline" and block.latex:
            return f"{self.math_delimiter_inline}{block.latex}{self.math_delimiter_inline}"
        
        elif block_type == "table" and block.table:
            return block.table.get("markdown", "")
        
        elif block_type == "figure" and block.figure_path:
            return f"![Figure]({block.figure_path})"
        
        elif block_type == "caption" and block.text:
            return f"*{block.text}*"
        
        elif block_type == "list" and block.text:
            # Convert to list format
            items = block.text.split('\n')
            return "\n".join(f"- {item.strip()}" for item in items if item.strip())
        
        elif block_type == "code" and block.text:
            return f"```\n{block.text}\n```"
        
        elif block.text:
            return block.text
        
        return ""
    
    def export_to_pdf(
        self,
        document: Any,
        output_path: Union[str, Path],
        engine: str = "auto"
    ) -> Optional[Path]:
        """
        Export document to PDF via Markdown conversion.
        
        Args:
            document: Document object
            output_path: Output PDF file path
            engine: Conversion engine ('auto', 'pandoc', 'weasyprint', 'md2pdf')
            
        Returns:
            Path to generated PDF, or None if conversion failed
        """
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        # Generate markdown first
        if hasattr(document, 'markdown') and document.markdown:
            markdown = document.markdown
        else:
            markdown = self._generate_markdown(document)
        
        # Create temp markdown file
        temp_md = output_path.with_suffix('.md.tmp')
        with open(temp_md, 'w', encoding='utf-8') as f:
            f.write(markdown)
        
        pdf_path = None
        
        # Try conversion engines in order of preference
        engines_to_try = []
        if engine == "auto":
            engines_to_try = ["pandoc", "weasyprint", "md2pdf"]
        else:
            engines_to_try = [engine]
        
        for eng in engines_to_try:
            try:
                if eng == "pandoc":
                    pdf_path = self._convert_with_pandoc(temp_md, output_path)
                elif eng == "weasyprint":
                    pdf_path = self._convert_with_weasyprint(temp_md, output_path, markdown)
                elif eng == "md2pdf":
                    pdf_path = self._convert_with_md2pdf(temp_md, output_path)
                
                if pdf_path and pdf_path.exists():
                    logger.info(f"Exported PDF using {eng}: {pdf_path}")
                    break
            except Exception as e:
                logger.debug(f"{eng} failed: {e}")
                continue
        
        # Cleanup temp file
        if temp_md.exists():
            temp_md.unlink()
        
        if not pdf_path or not pdf_path.exists():
            logger.error(
                "PDF conversion failed. Install one of:\n"
                "  - pandoc: https://pandoc.org/installing.html\n"
                "  - weasyprint: pip install weasyprint markdown\n"
                "  - md2pdf: pip install md2pdf"
            )
            return None
        
        return pdf_path
    
    def _convert_with_pandoc(self, md_path: Path, pdf_path: Path) -> Optional[Path]:
        """Convert markdown to PDF using pandoc."""
        result = subprocess.run(
            [
                "pandoc", str(md_path),
                "-o", str(pdf_path),
                "--pdf-engine=xelatex",
                "-V", "geometry:margin=1in"
            ],
            capture_output=True,
            text=True,
            timeout=120
        )
        if result.returncode == 0 and pdf_path.exists():
            return pdf_path
        return None
    
    def _convert_with_weasyprint(self, md_path: Path, pdf_path: Path, markdown: str) -> Optional[Path]:
        """Convert markdown to PDF using weasyprint."""
        try:
            import markdown as md_lib
            from weasyprint import HTML, CSS
            
            # Convert markdown to HTML
            html_content = md_lib.markdown(
                markdown,
                extensions=['tables', 'fenced_code', 'codehilite']
            )
            
            # Wrap in HTML document with styling
            full_html = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <meta charset="utf-8">
                <style>
                    body {{ font-family: 'Segoe UI', Arial, sans-serif; margin: 40px; line-height: 1.6; }}
                    h1 {{ color: #333; border-bottom: 2px solid #333; }}
                    h2 {{ color: #444; }}
                    table {{ border-collapse: collapse; width: 100%; margin: 20px 0; }}
                    th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
                    th {{ background-color: #f2f2f2; }}
                    code {{ background-color: #f4f4f4; padding: 2px 6px; border-radius: 3px; }}
                    pre {{ background-color: #f4f4f4; padding: 15px; overflow-x: auto; }}
                    blockquote {{ border-left: 4px solid #ddd; margin: 0; padding-left: 20px; color: #666; }}
                </style>
            </head>
            <body>
                {html_content}
            </body>
            </html>
            """
            
            HTML(string=full_html).write_pdf(str(pdf_path))
            return pdf_path
            
        except ImportError:
            raise ImportError("weasyprint or markdown not installed")
    
    def _convert_with_md2pdf(self, md_path: Path, pdf_path: Path) -> Optional[Path]:
        """Convert markdown to PDF using md2pdf."""
        try:
            from md2pdf.core import md2pdf as convert_md2pdf
            
            convert_md2pdf(
                pdf_path,
                md_file_path=str(md_path),
                css_file_path=None,
                base_url=None
            )
            return pdf_path if pdf_path.exists() else None
            
        except ImportError:
            raise ImportError("md2pdf not installed")


# ============================================================================
# DOCX Exporter
# ============================================================================

class DocxExporter:
    """Export document to DOCX format using python-docx."""
    
    def __init__(
        self,
        template_path: Optional[str] = None,
        equation_as_image: bool = True,
        image_width_inches: float = 6.0
    ):
        self.template_path = template_path
        self.equation_as_image = equation_as_image
        self.image_width_inches = image_width_inches
    
    def export(
        self,
        document: Any,
        output_path: Union[str, Path],
        figures_dir: Optional[Union[str, Path]] = None
    ) -> Path:
        """
        Export document to DOCX file.
        
        Args:
            document: Document object
            output_path: Output file path
            figures_dir: Directory containing figure images
            
        Returns:
            Path to the generated DOCX file
        """
        try:
            from docx import Document as DocxDocument
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise ImportError(
                "python-docx is required for DOCX export. "
                "Install with: pip install python-docx"
            )
        
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        # Create document from template or blank
        if self.template_path and Path(self.template_path).exists():
            doc = DocxDocument(self.template_path)
        else:
            doc = DocxDocument()
        
        # Use manifest if available
        if hasattr(document, 'docx_manifest') and document.docx_manifest:
            self._build_from_manifest(doc, document.docx_manifest, figures_dir)
        else:
            self._build_from_document(doc, document, figures_dir)
        
        # Save
        doc.save(str(output_path))
        logger.info(f"Exported DOCX to: {output_path}")
        
        return output_path
    
    def _build_from_manifest(
        self,
        doc: Any,
        manifest: Dict[str, Any],
        figures_dir: Optional[Path]
    ):
        """Build DOCX from manifest."""
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        elements = manifest.get("elements", [])
        
        for element in elements:
            elem_type = element.get("type")
            
            if elem_type == "heading":
                level = element.get("level", 1)
                text = element.get("text", "")
                if level == 0:
                    # Title
                    p = doc.add_heading(text, 0)
                else:
                    doc.add_heading(text, min(level, 9))
            
            elif elem_type == "paragraph":
                text = element.get("text", "")
                style = element.get("style")
                p = doc.add_paragraph(text)
                
                if style == "author":
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            elif elem_type == "equation":
                latex = element.get("latex", "")
                display = element.get("display", "block")
                
                if self.equation_as_image:
                    # TODO: Render LaTeX to image
                    # For now, add as text with note
                    p = doc.add_paragraph()
                    p.add_run(f"[Equation: {latex}]").italic = True
                    if display == "block":
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    # Add as plain text (OMML conversion would go here)
                    p = doc.add_paragraph(latex)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            elif elem_type == "table":
                rows = element.get("data", [])
                if rows:
                    self._add_table(doc, rows)
            
            elif elem_type == "image":
                path = element.get("path", "")
                if path and Path(path).exists():
                    doc.add_picture(path, width=Inches(self.image_width_inches))
                elif figures_dir:
                    # Try to find in figures dir
                    fig_path = Path(figures_dir) / Path(path).name
                    if fig_path.exists():
                        doc.add_picture(str(fig_path), width=Inches(self.image_width_inches))
    
    def _build_from_document(
        self,
        doc: Any,
        document: Any,
        figures_dir: Optional[Path]
    ):
        """Build DOCX from document object directly."""
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # Title
        if document.title:
            doc.add_heading(document.title, 0)
        
        # Authors
        if document.authors:
            p = doc.add_paragraph(document.authors)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Abstract
        if document.abstract:
            doc.add_heading("Abstract", level=1)
            doc.add_paragraph(document.abstract)
        
        # Process pages
        for page in document.pages:
            for block in sorted(page.blocks, key=lambda b: b.reading_order):
                self._add_block_to_docx(doc, block, figures_dir)
    
    def _add_block_to_docx(self, doc: Any, block: Any, figures_dir: Optional[Path]):
        """Add a block to the DOCX document."""
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        block_type = block.block_type if isinstance(block.block_type, str) else block.block_type.value
        
        if block_type == "heading" and block.text:
            doc.add_heading(block.text, level=2)
        
        elif block.text:
            doc.add_paragraph(block.text)
        
        elif block.latex:
            p = doc.add_paragraph()
            p.add_run(f"[Equation: {block.latex}]").italic = True
            if "block" in block_type:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        elif block.table:
            rows = block.table.get("struct", {}).get("rows", [])
            if rows:
                self._add_table(doc, rows)
        
        elif block.figure_path:
            path = Path(block.figure_path)
            if path.exists():
                doc.add_picture(str(path), width=Inches(self.image_width_inches))
    
    def _add_table(self, doc: Any, rows: List[List[str]]):
        """Add a table to the DOCX document."""
        if not rows:
            return
        
        num_rows = len(rows)
        num_cols = len(rows[0]) if rows else 0
        
        if num_cols == 0:
            return
        
        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Table Grid'
        
        for i, row_data in enumerate(rows):
            row = table.rows[i]
            for j, cell_text in enumerate(row_data):
                if j < len(row.cells):
                    row.cells[j].text = str(cell_text)


# ============================================================================
# LaTeX Exporter
# ============================================================================

class LatexExporter:
    """Export document to LaTeX format."""
    
    def __init__(
        self,
        document_class: str = "article",
        packages: Optional[List[str]] = None,
        font_size: str = "11pt"
    ):
        self.document_class = document_class
        self.packages = packages or [
            "amsmath", "amssymb", "graphicx", "booktabs",
            "hyperref", "geometry", "inputenc", "fontenc"
        ]
        self.font_size = font_size
    
    def export(
        self,
        document: Any,
        output_path: Union[str, Path],
        compile_pdf: bool = False
    ) -> Path:
        """
        Export document to LaTeX file.
        
        Args:
            document: Document object
            output_path: Output file path (.tex)
            compile_pdf: If True, compile to PDF using pdflatex
            
        Returns:
            Path to the generated LaTeX file (or PDF if compiled)
        """
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        latex = self._generate_latex(document)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(latex)
        
        logger.info(f"Exported LaTeX to: {output_path}")
        
        if compile_pdf:
            pdf_path = self._compile_pdf(output_path)
            return pdf_path
        
        return output_path
    
    def _generate_latex(self, document: Any) -> str:
        """Generate LaTeX source from document."""
        lines = []
        
        # Document class
        lines.append(f"\\documentclass[{self.font_size}]{{{self.document_class}}}")
        lines.append("")
        
        # Packages
        for pkg in self.packages:
            lines.append(f"\\usepackage{{{pkg}}}")
        lines.append("")
        
        # Geometry
        lines.append("\\geometry{margin=1in}")
        lines.append("")
        
        # Title, author
        if hasattr(document, 'title') and document.title:
            lines.append(f"\\title{{{self._escape_latex(document.title)}}}")
        if hasattr(document, 'authors') and document.authors:
            lines.append(f"\\author{{{self._escape_latex(document.authors)}}}")
        lines.append("\\date{}")
        lines.append("")
        
        # Begin document
        lines.append("\\begin{document}")
        lines.append("")
        
        if hasattr(document, 'title') and document.title:
            lines.append("\\maketitle")
            lines.append("")
        
        # Abstract
        if hasattr(document, 'abstract') and document.abstract:
            lines.append("\\begin{abstract}")
            lines.append(self._escape_latex(document.abstract))
            lines.append("\\end{abstract}")
            lines.append("")
        
        # Content
        for page in document.pages:
            for block in sorted(page.blocks, key=lambda b: b.reading_order):
                latex_block = self._block_to_latex(block)
                if latex_block:
                    lines.append(latex_block)
                    lines.append("")
        
        # End document
        lines.append("\\end{document}")
        
        return "\n".join(lines)
    
    def _block_to_latex(self, block: Any) -> str:
        """Convert a block to LaTeX."""
        block_type = block.block_type if isinstance(block.block_type, str) else block.block_type.value
        
        if block_type == "title":
            return ""  # Already handled in preamble
        
        elif block_type == "heading":
            if block.text:
                return f"\\section{{{self._escape_latex(block.text)}}}"
        
        elif block_type in ["paragraph", "abstract"]:
            if block.text:
                return self._escape_latex(block.text)
        
        elif block_type == "equation_block" and block.latex:
            return f"\\begin{{equation}}\n{block.latex}\n\\end{{equation}}"
        
        elif block_type == "equation_inline" and block.latex:
            return f"${block.latex}$"
        
        elif block_type == "table" and block.table:
            return self._table_to_latex(block.table)
        
        elif block_type == "figure" and block.figure_path:
            path = block.figure_path.replace('\\', '/')
            return (
                f"\\begin{{figure}}[h]\n"
                f"\\centering\n"
                f"\\includegraphics[width=0.8\\textwidth]{{{path}}}\n"
                f"\\end{{figure}}"
            )
        
        elif block_type == "caption" and block.text:
            return f"\\textit{{{self._escape_latex(block.text)}}}"
        
        elif block_type == "list" and block.text:
            items = block.text.split('\n')
            item_lines = "\n".join(
                f"\\item {self._escape_latex(item.strip())}"
                for item in items if item.strip()
            )
            return f"\\begin{{itemize}}\n{item_lines}\n\\end{{itemize}}"
        
        elif block.text:
            return self._escape_latex(block.text)
        
        return ""
    
    def _table_to_latex(self, table_data: Dict[str, Any]) -> str:
        """Convert table to LaTeX."""
        rows = table_data.get("struct", {}).get("rows", [])
        if not rows:
            return ""
        
        num_cols = len(rows[0]) if rows else 0
        col_spec = "|" + "c|" * num_cols
        
        lines = [
            "\\begin{table}[h]",
            "\\centering",
            f"\\begin{{tabular}}{{{col_spec}}}",
            "\\hline"
        ]
        
        for i, row in enumerate(rows):
            escaped_row = [self._escape_latex(str(cell)) for cell in row]
            lines.append(" & ".join(escaped_row) + " \\\\")
            lines.append("\\hline")
        
        lines.extend([
            "\\end{tabular}",
            "\\end{table}"
        ])
        
        return "\n".join(lines)
    
    def _escape_latex(self, text: str) -> str:
        """Escape special LaTeX characters."""
        if not text:
            return ""
        
        # Characters that need escaping
        special_chars = {
            '&': '\\&',
            '%': '\\%',
            '$': '\\$',
            '#': '\\#',
            '_': '\\_',
            '{': '\\{',
            '}': '\\}',
            '~': '\\textasciitilde{}',
            '^': '\\textasciicircum{}',
        }
        
        result = text
        for char, escaped in special_chars.items():
            result = result.replace(char, escaped)
        
        # Handle backslash specially (must be first)
        result = result.replace('\\', '\\textbackslash{}')
        
        return result
    
    def _compile_pdf(self, tex_path: Path) -> Path:
        """Compile LaTeX to PDF using pdflatex."""
        try:
            # Run pdflatex twice for references
            for _ in range(2):
                result = subprocess.run(
                    ["pdflatex", "-interaction=nonstopmode", str(tex_path.name)],
                    cwd=str(tex_path.parent),
                    capture_output=True,
                    text=True,
                    timeout=60
                )
            
            pdf_path = tex_path.with_suffix('.pdf')
            if pdf_path.exists():
                logger.info(f"Compiled PDF: {pdf_path}")
                return pdf_path
            else:
                logger.error(f"PDF compilation failed: {result.stderr}")
                return tex_path
                
        except subprocess.TimeoutExpired:
            logger.error("PDF compilation timed out")
            return tex_path
        except FileNotFoundError:
            logger.error(
                "pdflatex not found. Install LaTeX:\n"
                "  Windows: MiKTeX or TeX Live\n"
                "  macOS: MacTeX\n"
                "  Linux: sudo apt install texlive-latex-base"
            )
            return tex_path


# ============================================================================
# Multi-Format Exporter
# ============================================================================

class DocumentExporter:
    """Convenience class for exporting to multiple formats."""
    
    def __init__(
        self,
        output_dir: Union[str, Path],
        base_name: str = "document"
    ):
        self.output_dir = Path(output_dir)
        self.base_name = base_name
        
        self.markdown_exporter = MarkdownExporter()
        self.docx_exporter = DocxExporter()
        self.latex_exporter = LatexExporter()
    
    def export(
        self,
        document: Any,
        formats: List[str] = None
    ) -> Dict[str, Path]:
        """
        Export document to multiple formats.
        
        Args:
            document: Document object
            formats: List of formats ('markdown', 'docx', 'latex', 'pdf', 'pdf_md', 'all')
                     'pdf' uses LaTeX, 'pdf_md' uses Markdown conversion
            
        Returns:
            Dictionary mapping format to output path
        """
        if formats is None:
            formats = ["markdown", "docx", "latex"]
        
        if "all" in formats:
            formats = ["markdown", "docx", "latex", "pdf"]
        
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
        results = {}
        
        if "markdown" in formats:
            path = self.output_dir / f"{self.base_name}.md"
            results["markdown"] = self.markdown_exporter.export(document, path)
        
        if "docx" in formats:
            path = self.output_dir / f"{self.base_name}.docx"
            figures_dir = self.output_dir / "figures"
            results["docx"] = self.docx_exporter.export(
                document, path, figures_dir=figures_dir
            )
        
        # PDF via Markdown (simpler, no LaTeX required)
        if "pdf_md" in formats:
            path = self.output_dir / f"{self.base_name}.pdf"
            result = self.markdown_exporter.export_to_pdf(document, path)
            if result:
                results["pdf"] = result
        
        # PDF via LaTeX (better for math equations)
        if "latex" in formats or "pdf" in formats:
            path = self.output_dir / f"{self.base_name}.tex"
            compile_pdf = "pdf" in formats
            result = self.latex_exporter.export(
                document, path, compile_pdf=compile_pdf
            )
            
            if "latex" in formats:
                results["latex"] = path
            if "pdf" in formats and result.suffix == '.pdf':
                results["pdf"] = result
        
        return results


# ============================================================================
# Example Usage
# ============================================================================

if __name__ == "__main__":
    import sys
    import json
    
    # Simple test with mock document
    @dataclass
    class MockBlock:
        block_id: str
        block_type: str
        reading_order: int
        text: Optional[str] = None
        latex: Optional[str] = None
        table: Optional[Dict] = None
        figure_path: Optional[str] = None
    
    @dataclass
    class MockPage:
        page_number: int
        blocks: List[MockBlock]
    
    @dataclass
    class MockDocument:
        title: str
        authors: str
        abstract: str
        pages: List[MockPage]
        markdown: str = ""
        docx_manifest: Dict = None
    
    # Create mock document
    doc = MockDocument(
        title="Test Document",
        authors="Test Author",
        abstract="This is a test abstract.",
        pages=[
            MockPage(
                page_number=1,
                blocks=[
                    MockBlock("1", "heading", 0, text="Introduction"),
                    MockBlock("2", "paragraph", 1, text="This is a test paragraph."),
                    MockBlock("3", "equation_block", 2, latex="E = mc^2"),
                ]
            )
        ]
    )
    
    # Export
    output_dir = Path("./test_output")
    exporter = DocumentExporter(output_dir, "test_document")
    
    results = exporter.export(doc, formats=["markdown", "latex"])
    
    for fmt, path in results.items():
        print(f"{fmt}: {path}")


